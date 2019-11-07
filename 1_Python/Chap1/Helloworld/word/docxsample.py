import docx

doc =  docx.Document('TestData.docx')
print(doc.paragraphs)
len(doc.paragraphs)
print(doc.paragraphs[0].text)

for para in doc.paragraphs:
    print(para.text)