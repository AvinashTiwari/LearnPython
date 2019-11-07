import docx

doc =  docx.Document()
doc.add_paragraph("Hello World")
paraObject = doc.add_paragraph("This paraobject")
paraObject2 = doc.add_paragraph("some more paraobject")
paraObject.add_run("Adding more object")
doc.save("Newdoc.docx")
doc =  docx.Document("Newdoc.docx")
print(doc.paragraphs[0].style)
print(doc.paragraphs[0].runs)
doc.paragraphs[0].runs[0].bold = True
doc.paragraphs[0].runs[0].underline = True
doc.save("Newdoc.docx")
