import docx

doc =  docx.Document('TestData.docx')

print(doc.paragraphs)

print(doc.paragraphs[0].runs)
print(len(doc.paragraphs[0].runs))
print(doc.paragraphs[0].runs[1].text)